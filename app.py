import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import AffinityPropagation
from sklearn.decomposition import PCA
import matplotlib.pyplot as plt
from docx import Document
import os
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import re
import openai
import nltk

# Download NLTK data needed for text processing
nltk.download('stopwords')
nltk.download('punkt')

# Set OpenAI API key from environment variables for security
openai.api_key = "my api key"

# Function to read text from a .docx file
def read_docx(file_path):
    try:
        doc = Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        text = ' '.join(paragraphs)
        return text
    except Exception as e:
        st.error(f"Error reading file {file_path}: {str(e)}")
        return ""

# Function to remove common stopwords from text
def clean_stopwords(text: str) -> str:
    try:
        stopwords_list = ["a", "an", "and", "at", "but", "how", "in", "is", "on", "or", "the", "to", "what", "will"]
        tokens = word_tokenize(text)
        clean_tokens = [t for t in tokens if t.lower() not in stopwords_list]
        return " ".join(clean_tokens)
    except Exception as e:
        st.error(f"Error cleaning stopwords: {str(e)}")
        return text

# Function to clean up text by removing URLs, emails, and special characters
def clean_text(text):
    try:
        text = re.sub(r'http\S+', '', text)  # Remove URLs
        text = re.sub(r'\S+@\S+', '', text)  # Remove emails
        text = re.sub(r'[^a-zA-Z\s]', '', text)  # Remove non-alphabetic characters
        text = ' '.join(text.split())  # Remove extra spaces
        return text
    except Exception as e:
        st.error(f"Error cleaning text: {str(e)}")
        return text

# Function to cluster documents using Affinity Propagation algorithm
def cluster_documents(documents):
    try:
        vectorizer = TfidfVectorizer(stop_words='english')
        X = vectorizer.fit_transform(documents)  # Convert text to TF-IDF matrix
        model = AffinityPropagation()
        model.fit(X)  # Perform clustering
        labels = model.labels_  # Get the labels assigned to each document

        # Organize documents by cluster
        cluster_docs = {}
        for i in set(labels):
            cluster_docs[i] = [documents[j] for j in range(len(labels)) if labels[j] == i]
        return cluster_docs, X, labels
    except Exception as e:
        st.error(f"Error clustering documents: {str(e)}")
        return {}, None, None

# Function to visualize clusters using PCA to reduce dimensions
def visualize_clusters(X, labels):
    try:
        pca = PCA(n_components=2)
        scatter_plot_points = pca.fit_transform(X.toarray())  # Reduce dimensions to 2D
        colors = ['blue', 'green', 'red', 'purple', 'yellow', 'orange', 'black', 'brown', 'grey', 'pink']
        x_axis = [o[0] for o in scatter_plot_points]
        y_axis = [o[1] for o in scatter_plot_points]
        fig, ax = plt.subplots(figsize=(10, 5))

        # Plot points and color them based on their cluster
        for i in range(len(scatter_plot_points)):
            ax.scatter(x_axis[i], y_axis[i], c=colors[labels[i] % len(colors)], label=f'Cluster {labels[i]}')
        
        # Add legend and axis labels
        handles, labels = ax.get_legend_handles_labels()
        unique_labels = dict(zip(labels, handles))
        ax.legend(unique_labels.values(), unique_labels.keys())
        ax.set_title('PCA Visualization of Document Clusters')
        ax.set_xlabel('PCA Component 1')
        ax.set_ylabel('PCA Component 2')
        plt.tight_layout()
        st.pyplot(fig)
    except Exception as e:
        st.error(f"Error visualizing clusters: {str(e)}")

# Function to get all documents in a specified cluster
def access_files_of_cluster(cluster_docs, cluster_number):
    try:
        content = ""
        if cluster_number in cluster_docs:
            content = " ".join(cluster_docs[cluster_number])
        else:
            st.warning(f"Cluster {cluster_number} does not exist.")
        return content
    except Exception as e:
        st.error(f"Error accessing cluster files: {str(e)}")
        return ""

# Function to extract a company name from the text using OpenAI API
def extract_company_name(client, merged_text):
    try:
        tokens = word_tokenize(merged_text)
        first_500_tokens = ' '.join(tokens[:500])
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an AI assistant to extract a single, strong company name from the document."},
                {"role": "user", "content": f"Extract the company name from the given text:\n\n{first_500_tokens}"}
            ],
            max_tokens=60,
            temperature=0.6,
        )
        cluster_name = response.choices[0].message.content.strip()
        return cluster_name
    except Exception as e:
        st.error(f"Error extracting company name: {str(e)}")
        return ""

# Function to split text into chunks based on token count
def split_text_into_chunks(text, max_tokens=3000):
    try:
        tokens = word_tokenize(text)
        chunks = []
        if len(tokens) > max_tokens:
            middle_index = len(tokens) // 2
            chunks.append(' '.join(tokens[:middle_index]))
            chunks.append(' '.join(tokens[middle_index:]))
        else:
            chunks.append(' '.join(tokens))
        return chunks
    except Exception as e:
        st.error(f"Error splitting text into chunks: {str(e)}")
        return []

# Function to remove duplicates and summarize text using OpenAI API
def remove_duplicates_and_summarize(client, text, model):
    try:
        tokens = word_tokenize(text)
        if len(tokens) > 3000:
            # Split the text into smaller chunks if it exceeds token limit
            chunks = split_text_into_chunks(text, max_tokens=3000)
        else:
            chunks = [text]

        cleaned_chunks = []
        for chunk in chunks:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are an AI assistant that consolidates information,removing duplicates and maintaining all unique and relevant content."},
                    {"role": "user", "content": f"Your task is to process the following text by removing any
                      duplicate or semantically similar content while ensuring that no content repeats. The output
                      should be organized into clear, well-structured paragraphs. Each paragraph should present unique 
                     information without any duplication. The goal is to enhance clarity and conciseness, ensuring that
                      all unique insights are retained and presented in a coherent manner. Ensure that the text is
                      logically organized to facilitate better understanding and readability, and avoid repeating any 
                     content and give it as paragraph without repetition. :\n\n{chunk}"}
                ],
                max_tokens=3000,
                temperature=0.5,
            )
            cleaned_chunk = response.choices[0].message.content.strip()
            cleaned_chunks.append(cleaned_chunk)
        
        return ' '.join(cleaned_chunks)
    except Exception as e:
        st.error(f"Error removing duplicates and summarizing: {str(e)}")
        return ""

# Function to summarize text using OpenAI API
def summarize(client, text, model):
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an AI assistant that summarizes text to capture the main points and essential information."},
                {"role": "user", "content": f" Your task is to summarize the content; the content should be clean and summarized, and every detail should be mentioned without eliminating it. The output should be generally equal to the number of pages of input (atleat of 4 pages). Don't eliminate and repeat the content; the content should have side headings; if there are important things, they should be mentioned in points without repeating and eleminating original content. :\n\n{text}"}
            ],
            max_tokens=3000,
            temperature=0.5,
        )
        summary = response.choices[0].message.content.strip()
        return summary
    except Exception as e:
        st.error(f"Error summarizing text: {str(e)}")
        return ""

# Function to save text to a .docx file
def save_to_docx(text, output_path):
    try:
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)
        st.success(f"File saved successfully: {output_path}")
    except Exception as e:
        st.error(f"Error saving to docx: {str(e)}")

# Function to handle processing of a specific document cluster
def process_cluster(client, cluster_docs, cluster_number, output_dir_consolidated, output_dir_summarized, model, chunk_size=3000):
    try:
        # Get and merge text from the documents in the cluster
        merged_text = access_files_of_cluster(cluster_docs, cluster_number)

        # Extract a name for the company from the merged text
        company_name = extract_company_name(client, merged_text)

        # Split the text into manageable chunks
        chunks = split_text_into_chunks(merged_text, chunk_size)

        # Remove duplicates and summarize each chunk
        processed_chunks = [remove_duplicates_and_summarize(client, chunk, model) for chunk in chunks]
        final_text = " ".join(processed_chunks)

        # Save consolidated results
        consolidated_file_name = f"{company_name}_consolidated_output.docx" if company_name else f"cluster_{cluster_number}_Processed_Chunks.docx"
        consolidated_output_path = os.path.join(output_dir_consolidated, consolidated_file_name)
        save_to_docx(final_text, consolidated_output_path)

        # Create a final summary
        summary = summarize(client, final_text, model)

        # Save the summary to a .docx file
        summary_file_name = f"{company_name}_Summary_output.docx" if company_name else f"cluster_{cluster_number}_Summary.docx"
        summary_output_path = os.path.join(output_dir_summarized, summary_file_name)
        save_to_docx(summary, summary_output_path)
        
    except Exception as e:
        st.error(f"Error processing cluster {cluster_number}: {str(e)}")

# Streamlit app
st.title("Document Clustering and Summarization")

uploaded_files = st.file_uploader("Upload .docx files", type="docx", accept_multiple_files=True)
output_dir_consolidated = st.text_input("Consolidated Files Output Directory", "/path/to/consolidated")
output_dir_summarized = st.text_input("Summarized Files Output Directory", "/path/to/summarized")
process_button = st.button("Submit")

if process_button:
    if not uploaded_files:
        st.warning("Please upload at least one .docx file.")
    elif not output_dir_consolidated or not output_dir_summarized:
        st.warning("Please specify both output directories.")
    else:
        documents = [read_docx(file) for file in uploaded_files]
        st.write("Clustering documents...")
        cluster_docs, X, labels = cluster_documents(documents)
        
        if cluster_docs:
            st.write("Visualizing clusters...")
            visualize_clusters(X, labels)

            st.write("Processing clusters...")
            for cluster_number in cluster_docs.keys():
                process_cluster(openai, cluster_docs, cluster_number, output_dir_consolidated, output_dir_summarized, "gpt-4")
        else:
            st.warning("No clusters found.")
